# ============================================================
# TEXT FINGERPRINTING 
# ============================================================

import os
import re
import hashlib
import warnings
warnings.filterwarnings("ignore")

import zipfile
import numpy as np
from typing import Dict, Any

# ---------------- PDF ----------------
import pdfplumber
import PyPDF2

# ---------------- DOCX ----------------
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

# ---------------- NLP ----------------
import tensorflow as tf
import tensorflow_hub as hub
tf.get_logger().setLevel("ERROR")

# ---------------- WINDOWS DOC ----------------
try:
    import win32com.client
    import pythoncom
    DOC_SUPPORT = True
except ImportError:
    DOC_SUPPORT = False

# ---------------- NLP MODEL ----------------
try:
    NLP_MODEL = hub.load("https://tfhub.dev/google/universal-sentence-encoder/4")
except Exception:
    NLP_MODEL = None

# ============================================================
MIN_TOKEN_LEN = 50
MAX_CHARS = 8000
# ============================================================


# ============================================================
#  FILENAME WATERMARK VALIDATION
# ============================================================

def is_allowed_watermarked_docx(path: str) -> bool:
    """
    Allowed format:
    <name>_wm_<owner>.docx
    """
    fname = os.path.basename(path).lower()
    return (
        fname.endswith(".docx") and
        "_wm_" in fname and
        len(fname.split("_wm_")[-1].replace(".docx", "")) >= 3
    )


# ============================================================
#  ZIP VALIDATION (DOCX IS A ZIP)
# ============================================================

def is_valid_docx_zip(path: str) -> bool:
    try:
        with zipfile.ZipFile(path, "r") as z:
            return "[Content_Types].xml" in z.namelist()
    except Exception:
        return False


# ================= TXT =================
def extract_txt(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""


# ================= PDF =================
def extract_pdf(path: str) -> str:
    chunks = []

    try:
        with pdfplumber.open(path) as pdf:
            for p in pdf.pages:
                t = p.extract_text()
                if t:
                    chunks.append(t)
    except Exception:
        pass

    if not chunks:
        try:
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for p in reader.pages:
                    t = p.extract_text()
                    if t:
                        chunks.append(t)
        except Exception:
            pass

    return "\n".join(chunks)


# ============================================================
#  SAFE DOCX EXTRACTION 
# ============================================================

def extract_docx(path: str) -> str:
    if not is_valid_docx_zip(path):
        print(" INVALID DOCX ZIP STRUCTURE")
        return ""

    try:
        doc = Document(path)
        texts = []

        for p in doc.paragraphs:
            if p.text.strip():
                texts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text)

        return "\n".join(texts)

    except PackageNotFoundError:
        print(" DOCX PACKAGE ERROR")
        return ""
    except Exception as e:
        print(f" DOCX READ ERROR (python-docx): {e}")
        return ""


# ================= DOC (LEGACY WORD ONLY) =================
def extract_doc(path: str) -> str:
    if not DOC_SUPPORT:
        return ""

    try:
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

        doc = word.Documents.Open(os.path.abspath(path))
        text = doc.Content.Text

        doc.Close(False)
        word.Quit()
        pythoncom.CoUninitialize()
        return text.strip()

    except Exception as e:
        print(f" DOC READ ERROR: {e}")
        return ""


# ================= ROUTER =================
def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()

    if ext == ".txt":
        return extract_txt(path)
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".docx":
        return extract_docx(path)
    if ext == ".doc":
        return extract_doc(path)

    return ""


# ================= NORMALIZATION =================
def canonical_text(text: str) -> str:
    text = text.lower()
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^a-z0-9 ]", "", text)
    return text.strip()


# ================= NLP =================
def nlp_embedding(text: str) -> np.ndarray:
    if NLP_MODEL is None:
        return np.zeros(512, dtype=np.float32)

    vec = NLP_MODEL([text])[0].numpy()
    return vec / (np.linalg.norm(vec) + 1e-8)


# ================= FINGERPRINT =================
def fingerprint_text(file_path: str) -> Dict[str, Any]:
    raw = extract_text(file_path)

    status = "OK" if raw else "NO_TEXT"
    print(f"[{status}] {file_path} | chars={len(raw)} | words={len(raw.split())}")

    if not raw or len(raw.split()) < MIN_TOKEN_LEN:
        return {
            "hash": None,
            "nlp_vector": [],
            "raw_text": raw[:1000],
            "canonical_text": "",
            "normalized": False
        }

    canonical = canonical_text(raw)[:MAX_CHARS]

    return {
        "hash": hashlib.sha256(canonical.encode()).hexdigest(),
        "nlp_vector": nlp_embedding(canonical).tolist(),
        "raw_text": raw[:2000],
        "canonical_text": canonical,
        "normalized": True
    }
